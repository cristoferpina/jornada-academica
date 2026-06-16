import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"}, ...)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ["top", "left", "bottom", "right"]:
        if side in kwargs:
            side_prop = kwargs[side]
            tag = 'w:{}'.format(side)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, val in side_prop.items():
                element.set(qn('w:{}'.format(key)), str(val))

def generate_full_qa_plan():
    doc = Document()
    
    # --- 1. PORTADA ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\nPLAN DE PRUEBAS DE SOFTWARE")
    run.font.size = Pt(26)
    run.bold = True
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema de Gestión: 13va Jornada Académica y Cultural")
    run.font.size = Pt(20)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\nINGENIERÍA EN SISTEMAS COMPUTACIONALES")
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nPRESENTA:\n[NOMBRE DEL ALUMNO]\n\nMATERIA:\nPRUEBAS DE SOFTWARE\n\nDOCENTE:\n[NOMBRE DEL DOCENTE]\n\nFECHA:\n7 DE JUNIO DE 2026")
    run.font.size = Pt(12)

    doc.add_page_break()

    # --- 2. INTRODUCCIÓN ---
    doc.add_heading('2. Introducción', level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "En el desarrollo de sistemas críticos como el de la 13va Jornada Académica y Cultural, la calidad del software no es un atributo opcional, "
        "sino un requisito fundamental para el éxito institucional. El presente Plan de Pruebas define el marco de trabajo estratégico para evaluar "
        "la robustez, integridad y usabilidad de la plataforma. La importancia de las pruebas radica en su capacidad para mitigar riesgos operativos, "
        "asegurar la persistencia correcta de datos sensibles (ponentes, alumnos, recintos) y garantizar que los productos finales, como la papelería "
        "automatizada, cumplan con los estándares de impresión requeridos. Un sistema probado bajo metodologías rigurosas reduce los costos de mantenimiento "
        "y eleva la confianza del usuario final."
    )

    # --- 3. FUNDAMENTACIÓN TEÓRICA ---
    doc.add_heading('3. Fundamentación Teórica', level=1)
    
    theo = doc.add_paragraph()
    theo.add_run("3.1 Documentación del Diseño de Pruebas\n").bold = True
    theo.add_run(
        "La documentación del diseño de pruebas es el pilar de la trazabilidad en QA. Transforma requisitos abstractos en pasos ejecutables y resultados "
        "medibles. Según el estándar IEEE 829, un diseño de pruebas robusto permite que cualquier tester ejecute el mismo caso y obtenga el mismo resultado, "
        "eliminando la ambigüedad en la fase de control de calidad.\n\n"
    )
    
    theo.add_run("3.2 Verificación y Validación (V&V)\n").bold = True
    theo.add_run(
        "La Verificación se enfoca en responder: ¿Estamos construyendo el producto correctamente? Evalúa si el software cumple con las especificaciones técnicas "
        "y requerimientos definidos. La Validación, por otro lado, responde: ¿Estamos construyendo el producto correcto? Se centra en asegurar que el sistema "
        "satisface las necesidades reales de los administradores y alumnos de la jornada.\n\n"
    )
    
    theo.add_run("3.3 Técnicas de Caja Negra y Caja Blanca\n").bold = True
    theo.add_run(
        "Las técnicas de Caja Negra evalúan la funcionalidad desde la perspectiva del usuario, ignorando la estructura interna del código. Se basan en particiones "
        "de equivalencia y análisis de valores límite. La técnica de Caja Blanca (Caja de Cristal) requiere acceso al código fuente para evaluar rutas de control, "
        "bucles y condiciones lógicas, garantizando una cobertura estructural exhaustiva.\n\n"
    )

    theo.add_run("3.4 Niveles de Prueba: Unitarias, Integración y Regresión\n").bold = True
    theo.add_run(
        "Las Pruebas Unitarias validan el comportamiento de componentes aislados (ej. funciones de cálculo de aforo en Go). Las Pruebas de Integración verifican "
        "la comunicación entre el Frontend React y la API REST. Finalmente, las Pruebas de Regresión aseguran que nuevas modificaciones (como la actualización "
        "de colores institucionales) no corrompan funcionalidades existentes.\n\n"
    )

    theo.add_run("3.5 Pruebas de Carga y Estrés\n").bold = True
    theo.add_run(
        "Dada la naturaleza del evento, es vital evaluar cómo responde el sistema ante picos de tráfico simultáneo (ej. 500 alumnos consultando su matrícula al mismo tiempo). "
        "Las pruebas de carga determinan el comportamiento bajo condiciones normales, mientras que el estrés busca identificar el punto de quiebre del servidor.\n\n"
    )

    theo.add_run("3.6 Cobertura de Requisitos y Mejora Continua\n").bold = True
    theo.add_run(
        "La cobertura de requisitos es la métrica que indica qué porcentaje de las funciones prometidas han sido verificadas. Este proceso se retroalimenta en un "
        "ciclo de mejora continua, donde los hallazgos de cada fase de prueba guían el refinamiento del código y la prevención de futuros defectos."
    )

    doc.add_page_break()

    # --- 4. OBJETIVO ---
    doc.add_heading('4. Objetivo de las Pruebas', level=1)
    doc.add_paragraph("4.1 Objetivo General", style='Heading 2')
    doc.add_paragraph("Validar integralmente el sistema de gestión de la 13va Jornada Académica, asegurando que todos los flujos administrativos y de consulta operen sin fallos críticos.")
    doc.add_paragraph("4.2 Objetivos Específicos", style='Heading 2')
    spec_list = [
        "Verificar la persistencia inmediata de la asignación de recintos en la base de datos.",
        "Validar la generación de documentos PDF (papelería) con fidelidad visual del 100%.",
        "Asegurar la integridad del padrón de alumnos mediante validaciones de concurrencia y duplicidad.",
        "Evaluar la respuesta del backend ante peticiones autenticadas y no autorizadas."
    ]
    for item in spec_list:
        doc.add_paragraph(item, style='List Bullet')

    # --- 5. ALCANCE ---
    doc.add_heading('5. Alcance del Plan de Pruebas', level=1)
    doc.add_paragraph("En Alcance:", style='Heading 2')
    scope_in = ["Módulo de Ponentes (CRUD y carga de archivos)", "Módulo de Recintos (Registro y estado)", "Gestión de Alumnos (Vinculación)", "Generador de Papelería (Exportación PDF)", "Seguridad (JWT y Roles)"]
    for item in scope_in: doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph("Fuera de Alcance:", style='Heading 2')
    scope_out = ["Despliegue en servidores de producción finales", "Compatibilidad con navegadores obsoletos (IE11)", "Pruebas de penetración de red externas"]
    for item in scope_out: doc.add_paragraph(item, style='List Bullet')

    # --- 6. MÓDULOS A EVALUAR ---
    doc.add_heading('6. Módulos o Funcionalidades a Evaluar', level=1)
    table_m = doc.add_table(rows=1, cols=4)
    table_m.style = 'Table Grid'
    hdr = table_m.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'ID', 'Módulo', 'Descripción', 'Prioridad'
    
    mod_data = [
        ('M-01', 'Seguridad', 'Autenticación JWT y control de acceso por roles.', 'Alta'),
        ('M-02', 'Ponentes', 'Gestión de perfiles, fotos y asignación de temas.', 'Alta'),
        ('M-03', 'Recintos', 'Control de espacios físicos y disponibilidad.', 'Media'),
        ('M-04', 'Alumnos', 'Padrón institucional y vinculación de cuentas.', 'Alta'),
        ('M-05', 'Papelería', 'Generación de panfletos, portanombres y etiquetas.', 'Crítica')
    ]
    for r in mod_data:
        row = table_m.add_row().cells
        for i, val in enumerate(r): row[i].text = val

    # --- 7. ESTRATEGIA ---
    doc.add_heading('7. Estrategia de Pruebas', level=1)
    strat = doc.add_paragraph()
    strat.add_run("Estrategia de Caja Negra:").bold = True
    strat.add_run(" Se realizarán pruebas funcionales basadas en la interfaz de usuario para validar que el ingreso de datos produzca la salida esperada.\n")
    strat.add_run("Estrategia de Caja Blanca:").bold = True
    strat.add_run(" Revisión de los controladores en Go para asegurar que el manejo de errores SQL no exponga vulnerabilidades.\n")
    strat.add_run("Pruebas de Integración:").bold = True
    strat.add_run(" Validación de los endpoints de la API asegurando que el JSON de respuesta sea consumible por el Frontend.\n")
    strat.add_run("Pruebas de Regresión:").bold = True
    strat.add_run(" Ejecución de suites de pruebas automáticas tras cada actualización de estilos o librerías.")

    doc.add_page_break()

    # --- 8. CASOS DE PRUEBA (15 Casos) ---
    doc.add_heading('8. Casos de Prueba', level=1)
    
    def add_tc(id_tc, name, desc, req, inputs, pre, proc, expected):
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "ID: " + id_tc
        table.rows[0].cells[1].text = "Nombre: " + name
        table.rows[1].cells[0].merge(table.rows[1].cells[1]).text = "Descripción: " + desc
        table.rows[2].cells[0].text = "Requisito: " + req
        table.rows[2].cells[1].text = "Estado: Pendiente"
        table.rows[3].cells[0].merge(table.rows[3].cells[1]).text = "Entradas: " + inputs
        table.rows[4].cells[0].merge(table.rows[4].cells[1]).text = "Precondiciones: " + pre
        table.rows[5].cells[0].merge(table.rows[5].cells[1]).text = "Procedimiento: " + proc
        table.rows[6].cells[0].merge(table.rows[6].cells[1]).text = "Resultado Esperado: " + expected
        doc.add_paragraph("")

    # --- LISTA DE 15 CASOS ---
    # Ponentes
    add_tc("CP-PON-01", "Registro Positivo de Ponente", "Validar registro con todos los datos válidos", "RF-02", "Nombre, Grado, Foto (JPG)", "Admin logueado", "1. Llenar campos. 2. Subir imagen. 3. Click Guardar.", "Mensaje de éxito y registro en tabla.")
    add_tc("CP-PON-02", "Registro Negativo - Imagen Inválida", "Validar rechazo de archivos no soportados", "RF-02", "Archivo .exe", "Admin logueado", "1. Seleccionar archivo .exe en campo foto.", "Sistema bloquea la carga o muestra error de formato.")
    
    # Recintos
    add_tc("CP-REC-03", "Asignación Automática", "Validar registro inmediato al seleccionar", "RF-03", "Selección en Combobox", "Ponente en modo edición", "1. Cambiar recinto. 2. Refrescar página.", "El recinto persiste sin dar click en 'Actualizar'.")
    add_tc("CP-REC-04", "Registro Fallido - Nombre Duplicado", "Evitar nombres de recintos iguales", "RF-03", "Nombre: 'Auditorio A'", "Auditorio A ya existe", "1. Intentar crear nuevo recinto con nombre existente.", "Sistema muestra advertencia de duplicidad.")
    
    # Alumnos
    add_tc("CP-ALU-05", "Búsqueda en Tiempo Real", "Validar filtrado dinámico", "RF-04", "Matrícula: '2209'", "Existen alumnos con esa matrícula", "1. Escribir '2209' en el buscador.", "La tabla se reduce a los alumnos que coinciden.")
    add_tc("CP-ALU-06", "Registro Alumno Positivo", "Alta de nuevo estudiante", "RF-04", "Matrícula, Nombre, Carrera", "Admin logueado", "1. Formulario > Crear.", "Alumno aparece en lista de 'Pendientes'.")
    add_tc("CP-ALU-07", "Error Matrícula Incompleta", "Validar longitud de matrícula", "RF-04", "Matrícula: '123'", "Admin logueado", "1. Ingresar matrícula corta. 2. Guardar.", "Error: La matrícula debe tener 8 dígitos.")
    
    # Papelería
    add_tc("CP-PAP-08", "Generación de Panfleto", "Validar exportación PDF tríptico", "RF-05", "Click en 'Panfleto'", "Ponente registrado", "1. Click botón panfleto.", "Descarga de PDF con 3 columnas y datos correctos.")
    add_tc("CP-PAP-09", "Etiqueta Botella - Datos correctos", "Validar info en etiqueta", "RF-05", "Click en 'Etiqueta'", "Ponente con nombre largo", "1. Generar etiqueta.", "Nombre ajustado al tamaño de la etiqueta (20x5cm).")
    add_tc("CP-PAP-10", "Portanombre - Efecto Espejo", "Validar impresión para doblado", "RF-05", "Click en 'Portanombre'", "Admin en Papelería", "1. Generar PDF.", "Una cara aparece invertida para correcto doblado carpa.")
    
    # Seguridad
    add_tc("CP-SEG-11", "Login Exitoso", "Acceso de superusuario", "RF-01", "Email/Pass válidos", "Usuario registrado", "1. Ingresar credenciales.", "Redirección al Dashboard.")
    add_tc("CP-SEG-12", "Bloqueo por Credenciales Erróneas", "Validar seguridad de acceso", "RF-01", "Email correcto / Pass incorrecto", "N/A", "1. Intentar login.", "Mensaje: 'Credenciales inválidas'. No permite acceso.")
    add_tc("CP-SEG-13", "Expiración de Sesión", "Validar seguridad temporal", "RF-01", "Token expirado", "Sesión inactiva 24h", "1. Intentar acción.", "Redirección automática al Login.")
    
    # Generales
    add_tc("CP-GEN-14", "Responsive Design - Mobile", "Validar vista en dispositivos móviles", "RNF-01", "Vista iPhone 12 screen", "Acceso desde móvil", "1. Abrir dashboard en móvil.", "Sidebar se oculta o adapta; botones legibles.")
    add_tc("CP-GEN-15", "Carga de Imágenes Grandes", "Validar límite de MB", "RF-02", "Foto de 20MB", "Campo foto perfil", "1. Subir archivo pesado.", "Error: 'El archivo excede los 5MB permitidos'.")

    # --- 9. DATOS DE PRUEBA ---
    doc.add_heading('9. Datos de Prueba', level=1)
    doc.add_paragraph("Para garantizar la cobertura, se utilizarán los siguientes sets de datos:")
    table_d = doc.add_table(rows=1, cols=3)
    table_d.style = 'Table Grid'
    h = table_d.rows[0].cells
    h[0].text, h[1].text, h[2].text = 'Tipo', 'Ejemplo Válido', 'Ejemplo Inválido / Límite'
    d_rows = [
        ('Matrículas', '20230001, 21094556', 'ABCD-123 (Formato), 1 (Límite inferior)'),
        ('Imágenes', 'foto.jpg, logo.png', 'documento.pdf, script.sh (Inyección)'),
        ('Fechas', '2026-06-15', '2020-01-01 (Pasado), 2099-12-31 (Futuro lejano)'),
        ('Capacidad', '50, 100', '-5 (Negativo), 10000 (Exceso)')
    ]
    for r in d_rows:
        row = table_d.add_row().cells
        row[0].text, row[1].text, row[2].text = r

    # --- 10. RESULTADOS ESPERADOS ---
    doc.add_heading('10. Resultados Esperados', level=1)
    doc.add_paragraph(
        "Se espera que el sistema mantenga una estabilidad operativa del 99.9%. "
        "Específicamente: 1. La base de datos no debe presentar registros huérfanos. "
        "2. Los PDFs deben generarse en menos de 3 segundos. "
        "3. La interfaz debe ser intuitiva, requiriendo máximo 3 clics para cualquier acción administrativa principal."
    )

    # --- 11. CRITERIOS DE ACEPTACIÓN ---
    doc.add_heading('11. Criterios de Aceptación', level=1)
    criteria = [
        "100% de los casos de prueba clasificados como 'Críticos' deben estar en estado: Aprobado.",
        "Ausencia total de defectos de Severidad 1 (Bloqueante).",
        "Menos de 2 defectos de Severidad 2 (Alta) con plan de corrección a corto plazo.",
        "Carga inicial del Dashboard en menos de 1.5 segundos en red institucional."
    ]
    for c in criteria: doc.add_paragraph(c, style='List Bullet')

    # --- 12. RIESGOS ---
    doc.add_heading('12. Riesgos Detectados', level=1)
    table_r = doc.add_table(rows=1, cols=4)
    table_r.style = 'Table Grid'
    h = table_r.rows[0].cells
    h[0].text, h[1].text, h[2].text, h[3].text = 'Riesgo', 'Probabilidad', 'Impacto', 'Mitigación'
    r_data = [
        ('Pérdida de sesión en media carga', 'Baja', 'Media', 'Persistencia en LocalStorage.'),
        ('PDF con caracteres corruptos', 'Media', 'Alta', 'Uso de fuentes estándar (Helvetica).'),
        ('Saturación de base de datos', 'Baja', 'Muy Alta', 'Optimización de índices y pools.'),
        ('Incompatibilidad móvil', 'Baja', 'Media', 'Pruebas con Chrome DevTools Responsive.')
    ]
    for r in r_data:
        row = table_r.add_row().cells
        for i, v in enumerate(r): row[i].text = v

    # --- 13. CRONOGRAMA ---
    doc.add_heading('13. Cronograma de Actividades', level=1)
    table_c = doc.add_table(rows=1, cols=2)
    table_c.style = 'Table Grid'
    h = table_c.rows[0].cells
    h[0].text, h[1].text = 'Semana', 'Actividad'
    c_data = [
        ('Semana 1', 'Planeación estratégica y análisis de requisitos.'),
        ('Semana 2', 'Diseño de casos de prueba y preparación de datos.'),
        ('Semana 3', 'Ejecución de pruebas funcionales y reporte de errores.'),
        ('Semana 4', 'Pruebas de regresión, carga y validación final.')
    ]
    for r in c_data:
        row = table_c.add_row().cells
        row[0].text, row[1].text = r

    # --- 14. RECURSOS ---
    doc.add_heading('14. Recursos Necesarios', level=1)
    rec = doc.add_paragraph()
    rec.add_run("Hardware:").bold = True
    rec.add_run(" Laptop con 8GB RAM mín., Servidor local para BD.\n")
    rec.add_run("Software:").bold = True
    rec.add_run(" Node.js, Go Compiler, PostgreSQL 14.\n")
    rec.add_run("Herramientas QA:").bold = True
    rec.add_run(" Postman (APIs), Chrome DevTools, JMeter (Carga).")

    # --- 15. MÉTRICAS ---
    doc.add_heading('15. Métricas e Indicadores', level=1)
    metrics = [
        "Porcentaje de Cobertura: (Requisitos probados / Requisitos totales) * 100",
        "Tasa de Aprobación: (Casos exitosos / Casos ejecutados) * 100",
        "Densidad de Defectos: (N° Defectos / KLOC o Módulo)",
        "Tiempo de Resolución: Promedio de horas desde reporte a corrección."
    ]
    for m in metrics: doc.add_paragraph(m, style='List Bullet')

    # --- 16. CONCLUSIONES ---
    doc.add_heading('16. Conclusiones', level=1)
    doc.add_paragraph(
        "El proceso de pruebas es la única garantía científica de que el sistema responderá adecuadamente ante la presión real del evento. "
        "La implementación de este plan asegura no solo que el software funcione, sino que la inversión institucional en tecnología se traduzca "
        "en eficiencia administrativa. La calidad no es un destino, sino un viaje de mejora continua fundamentado en la documentación rigurosa."
    )

    doc.save('Plan_de_Pruebas_Senior_QA_Jornada.docx')

if __name__ == '__main__':
    generate_full_qa_plan()
