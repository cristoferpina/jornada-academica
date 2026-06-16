from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_plan():
    doc = Document()
    
    # Estilo de Título
    title = doc.add_heading('Plan de Pruebas de Software', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Proyecto: Sistema de Gestión - 13va Jornada Académica y Cultural')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Objetivo
    doc.add_heading('1. Objetivo de las Pruebas', level=1)
    doc.add_paragraph(
        "Garantizar la calidad, integridad y correcto funcionamiento de la plataforma administrativa de la 13va Jornada Académica. "
        "El enfoque principal es la detección temprana de fallos mediante procesos de Verificación (cumplimiento técnico) "
        "y Validación (satisfacción de necesidades del usuario), asegurando que cada requisito funcional sea una evidencia verificable."
    )

    # 2. Alcance
    doc.add_heading('2. Alcance del Plan de Pruebas', level=1)
    doc.add_paragraph(
        "El alcance abarca los módulos críticos del backend (Go) y frontend (React), específicamente: "
        "Gestión de Ponentes, Asignación de Recintos, Registro de Alumnos y Generación de Material de Papelería en PDF."
    )

    # 3. Módulos a evaluar
    doc.add_heading('3. Módulos o Funcionalidades a Evaluar', level=1)
    table_mod = doc.add_table(rows=1, cols=2)
    table_mod.style = 'Table Grid'
    hdr_cells = table_mod.rows[0].cells
    hdr_cells[0].text = 'Módulo'
    hdr_cells[1].text = 'Descripción'
    
    modules = [
        ('Ponentes', 'Gestión de perfiles, fotos y carga de logos institucionales.'),
        ('Recintos', 'Asignación inmediata de espacios físicos y visualización de detalles.'),
        ('Alumnos', 'Control de padrón, vinculación de cuentas y búsqueda por matrícula.'),
        ('Papelería', 'Generación dinámica de PDFs (Panfletos, Portanombres, Etiquetas).')
    ]
    for m, d in modules:
        row_cells = table_mod.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = d

    # 4. Casos de Prueba
    doc.add_heading('4. Casos de Prueba', level=1)
    doc.add_paragraph("Se aplican estrategias de Caja Negra para validación funcional y de usabilidad.")
    
    # Caso 1
    doc.add_heading('CP-REC-01: Registro inmediato de Recinto', level=2)
    doc.add_paragraph("Descripción: Validar que al seleccionar un recinto en el combobox, el cambio se registre automáticamente en la BD.")
    doc.add_paragraph("Entradas: Selección de 'Auditorio Principal' en el menú desplegable de un ponente en modo edición.")
    doc.add_paragraph("Procedimiento:")
    doc.add_paragraph("1. Iniciar sesión como Administrador.\n2. Ir a Ponentes > Editar.\n3. Cambiar el valor del Combobox 'Recinto'.\n4. Verificar llamada a la API (PUT /speakers/:id).")
    doc.add_paragraph("Resultado Esperado: El recinto se actualiza en la base de datos sin necesidad de guardar manualmente el formulario.")

    # Caso 2
    doc.add_heading('CP-ALU-02: Registro de Alumno (Caso Negativo)', level=2)
    doc.add_paragraph("Descripción: Validar el manejo de matrículas duplicadas.")
    doc.add_paragraph("Entradas: Datos de alumno con una matrícula ya existente en el sistema.")
    doc.add_paragraph("Procedimiento:")
    doc.add_paragraph("1. Ir a Gestión de Alumnos.\n2. Clic en 'Registrar Alumno'.\n3. Ingresar una matrícula existente.\n4. Intentar guardar.")
    doc.add_paragraph("Resultado Esperado: El sistema debe mostrar un mensaje de error claro indicando que la matrícula ya está registrada.")

    # 5. Datos de Prueba
    doc.add_heading('5. Datos de Prueba', level=1)
    doc.add_paragraph("- Matrículas reales (ej. 22090001) y ficticias.\n- Imágenes en formato .jpg y .png para perfiles.\n- Nombres con caracteres especiales (acentos/eñes).")

    # 6. Resultados Esperados y Criterios de Aceptación
    doc.add_heading('6. Criterios de Aceptación', level=1)
    doc.add_paragraph("- 100% de los casos positivos ejecutados con éxito.\n- Los PDFs generados deben ser legibles y con los colores institucionales (#80ba26).\n- El tiempo de respuesta de la API no debe exceder los 2 segundos.")

    # 7. Riesgos Detectados
    doc.add_heading('7. Riesgos Detectados', level=1)
    doc.add_paragraph("- Pérdida de integridad de datos por concurrencia.\n- Incompatibilidad de fuentes OTF en visores de PDF antiguos (Mitigado usando Helvetica).")

    # 8. Cronograma
    doc.add_heading('8. Cronograma de Actividades', level=1)
    doc.add_paragraph("- Fase 1: Definición y Especificaciones (Día 1).\n- Fase 2: Desarrollo de Reactivos (Día 2).\n- Fase 3: Ejecución y Piloteo (Día 3).\n- Fase 4: Análisis de Resultados (Día 4).")

    # 9. Recursos
    doc.add_heading('9. Recursos Necesarios', level=1)
    doc.add_paragraph("- Entorno de desarrollo local (Frontend React / Backend Go).\n- Navegador Chrome/Firefox última versión.\n- Base de Datos PostgreSQL.")

    doc.save('Plan_de_Pruebas_Jornada_Academica.docx')

if __name__ == '__main__':
    generate_plan()
